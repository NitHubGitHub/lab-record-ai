from fastapi import FastAPI
from pydantic import BaseModel
import json
from bs4 import BeautifulSoup
from docx.shared import Pt

from docx.enum.style import WD_STYLE_TYPE
from docx.shared import RGBColor

from docx import Document
import uuid
import os

from prompt_templates import build_rules
from fastapi.responses import FileResponse
from fastapi.staticfiles import StaticFiles

from docx.shared import RGBColor

from io import BytesIO
from fastapi.responses import StreamingResponse


from ai_provider import (
    validate_experiment,
    generate_lab_record
)


app = FastAPI()



class ValidationRequest(BaseModel):
    course: str
    experiment: str


class GenerateRequest(BaseModel):
    course: str
    experiment: str
    sections: list[str]

app.mount("/static", StaticFiles(directory="static"), name="static")




@app.get("/")
def home():
    return FileResponse("static/index.html")

@app.get("/health")
def health():
    return {"status": "ok"}


@app.post("/validate")
def validate(data: ValidationRequest):

    prompt = f"""
You are an academic laboratory experiment validator.

Course:
{data.course}

Experiment:
{data.experiment}

Tasks:

1. Correct spelling mistakes in both fields.
2. Determine whether the experiment is a legitimate academic laboratory experiment, practical exercise, simulation, engineering study, or scientific investigation.
3. Identify the most appropriate domain.
4. Return ONLY valid JSON.

VALID EXPERIMENTS include:

- Programming laboratory exercises
- Data structures experiments
- Electronics circuits
- Sensor experiments
- DBMS practicals
- Physics experiments
- Chemistry experiments
- Mechanical engineering experiments
- Civil engineering experiments
- Biology laboratory activities
- Simulations
- Measurements
- Scientific investigations
- Testing and analysis exercises

IMPORTANT:

A topic is NOT valid simply because it belongs to an academic subject.

The experiment must involve:

- implementation
- testing
- measurement
- observation
- analysis
- simulation
- experimentation
- investigation

INVALID examples:

Course: F1
Experiment: Best F1 Drivers

Course: Football
Experiment: Greatest Players of All Time

Course: Cars
Experiment: Best Luxury Cars

Course: Programming
Experiment: Top Programming Languages

Course: Physics
Experiment: Most Famous Physicists

These are rankings, comparisons, opinions, discussions, or general topics.
They are NOT laboratory experiments.

Domains:

Programming
Data Structures
DBMS
Sensors
Electronics
Electrical
Physics
Chemistry
Mechanical
Civil
Biology
Other

Return ONLY valid JSON.

For valid experiments, also return a field called
"recommended_sections".

IMPORTANT:

You must ONLY choose section names from the list below.

Do NOT invent new section names.
Do NOT modify section names.
Do NOT combine section names.

Allowed sections:

Aim
Requirements
Theory
Diagram (not image, but made using lines)
Formula
Algorithm
Flowchart
Procedure
Program
Source Code
Observation
Observation Table
Calculation
Input
Output
Result
Conclusion
Applications
Precautions
Viva Questions
Learning Outcomes
References

Return recommended_sections as a JSON array.

Example:

{{
  "valid": true,
  "domain": "Programming",
  "corrected_course": "Programming in C",
  "corrected_experiment": "Matrix Addition Using Functions",
  "recommended_sections": [
    "Aim",
    "Theory",
    "Algorithm",
    "Program",
    "Output",
    "Result"
  ],
  "reason": ""
}}
"""

    try:

        response_text = validate_experiment(prompt)

    except Exception:

        return {
            "valid": False,
            "domain": "System",
            "corrected_course": data.course,
            "corrected_experiment": data.experiment,
            "reason": "Validation service temporarily unavailable, please try again later"
        }

    try:

        cleaned = response_text.strip()

        if cleaned.startswith("```json"):
            cleaned = cleaned.replace("```json", "", 1)

        if cleaned.endswith("```"):
            cleaned = cleaned[:-3]

        cleaned = cleaned.strip()

        result = json.loads(cleaned)

        return result

    except Exception as e:

        return {
            "error": str(e),
            "raw_response": response_text
        }

@app.post("/generate")
def generate(data: GenerateRequest):

    sections_text = "\n".join(data.sections)

    rules = build_rules(data.sections)

    prompt = f"""
You are an expert engineering laboratory record writer.

Course:
{data.course}

Experiment:
{data.experiment}

Sections Requested:
{sections_text}

{rules}

Generate only the requested sections.

The output must be suitable for direct writing into an engineering laboratory record notebook.

Use proper section headings.

Follow all formatting and length requirements exactly.

Do not provide explanations outside the requested sections.
"""

    content = generate_lab_record(prompt)

    return {
        "content": content
    }

@app.post("/download-docx")
async def download_docx(data: dict):

    content = data["content"]

    doc = Document()




    heading_style = doc.styles["Heading 1"]

    heading_style.font.color.rgb = RGBColor(
        96,
        165,
        250
    )

    title = doc.add_heading(level=0)

    title_run = title.add_run(
        "Laboratory Record"
    )

    title_run.bold = True

    title_run.font.size = Pt(28)
    title_run.font.color.rgb = RGBColor(
        30,
        58,
        138
    )

    for line in content.split("\n"):

        line = line.strip()

        if not line:
            continue

        if line.startswith("##"):

            heading = line.replace(
                "##",
                ""
            ).strip()

            h = doc.add_heading(
                heading,
                level=1
            )

            h.paragraph_format.space_before = Pt(16)

            h.paragraph_format.space_after = Pt(12)

            continue

        paragraph = doc.add_paragraph()

        paragraph.paragraph_format.space_after = Pt(6)

        run = paragraph.add_run(line)

        run.font.size = Pt(14)

    buffer = BytesIO()

    doc.save(buffer)

    buffer.seek(0)

    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition":
            "attachment; filename=Lab_Record.docx"
        }
    )